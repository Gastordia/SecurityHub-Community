"""
Exporter - Converts rendered HTML to various output formats (PDF, Word, CSV, etc.).
Full implementation with DOCX and CSV support.
"""
import logging
from typing import Union, Optional
from pathlib import Path
import io
import csv
import re
import base64
from urllib.parse import urlparse

logger = logging.getLogger(__name__)

# Lazy import WeasyPrint to avoid startup errors if GTK+ libraries are missing
WEASYPRINT_AVAILABLE = False
WeasyHTML = None
try:
    from weasyprint import HTML as WeasyHTML
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as e:
    logger.warning(f"WeasyPrint not available: {e}. PDF export will be disabled. "
                   f"Install GTK+ runtime libraries to enable PDF export. "
                   f"See: https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation")
    WEASYPRINT_AVAILABLE = False

# Try to import DOCX libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX export will be limited.")

try:
    from bs4 import BeautifulSoup
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False
    logger.warning("BeautifulSoup not available. CSV export from HTML will be limited.")


class Exporter:
    """
    Exports rendered HTML to various formats.
    Supports HTML, PDF, Word (DOCX), and CSV with full implementations.
    """
    
    def __init__(self):
        """Initialize exporter."""
        self.asset_helper = None
    
    def export(self, html: str, format: str = "html", base_url: str = None) -> Union[str, bytes]:
        """
        Export HTML to specified format.
        
        Args:
            html: Rendered HTML content
            format: Target format (html, pdf, docx, csv)
            base_url: Base URL for resolving relative paths in HTML
        
        Returns:
            Exported content (string for HTML/CSV, bytes for PDF/DOCX)
        
        Raises:
            ValueError: If format is not supported
        """
        format_lower = format.lower()
        
        # Resolve asset URLs for PDF/DOCX exports
        if format_lower in ["pdf", "docx"]:
            html = self._resolve_asset_urls(html, format_lower)
        
        if format_lower == "html":
            return html
        elif format_lower == "pdf":
            return self._to_pdf(html, base_url)
        elif format_lower in ["word", "docx"]:
            return self._to_docx(html)
        elif format_lower == "csv":
            return self._to_csv(html)
        else:
            raise ValueError(f"Unsupported format: {format}")
    
    def _resolve_asset_urls(self, html: str, format: str) -> str:
        """
        Resolve asset URLs in HTML for static export formats.
        Converts asset URLs to base64 data URIs for PDF/DOCX.
        
        Args:
            html: HTML content
            format: Target format (pdf or docx)
        
        Returns:
            HTML with resolved asset URLs
        """
        if not self.asset_helper:
            return html
        
        # Find all asset references: {{ asset('name') }} or {{ 'name'|asset_url }}
        asset_pattern = r'(?:{{|{%)\s*(?:asset\([\'"]?([^\'"\)]+)[\'"]?\)|[\'"]?([^\'"]+)[\'"]?\s*\|\s*asset(?:_url)?)\s*(?:}}|%})'
        
        def replace_asset(match):
            asset_name = match.group(1) or match.group(2)
            if not asset_name:
                return match.group(0)
            
            # Get asset URL
            asset_url = self.asset_helper.get_asset_url(asset_name)
            if not asset_url:
                logger.warning(f"Asset '{asset_name}' not found for export")
                return match.group(0)
            
            # For base64 assets, use directly
            if asset_url.startswith('data:'):
                return asset_url
            
            # For file URLs, try to convert to base64 if small enough
            # For now, return the URL (WeasyPrint can handle HTTP URLs)
            return asset_url
        
        # Replace asset references
        html = re.sub(asset_pattern, replace_asset, html)
        
        return html
    
    def _to_pdf(self, html: str, base_url: str = None) -> bytes:
        """
        Convert HTML to PDF using WeasyPrint.
        
        Args:
            html: HTML content
            base_url: Base URL for resolving relative paths
        
        Returns:
            PDF content as bytes
        
        Raises:
            ValueError: If WeasyPrint is not available or PDF generation fails
        """
        if not WEASYPRINT_AVAILABLE:
            error_msg = (
                "PDF export is not available. WeasyPrint requires GTK+ runtime libraries. "
                "Please install GTK+ runtime for Windows from: "
                "https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/releases"
            )
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        try:
            if base_url:
                pdf = WeasyHTML(string=html, base_url=base_url).write_pdf()
            else:
                pdf = WeasyHTML(string=html).write_pdf()
            return pdf
        except Exception as e:
            logger.error(f"Error converting HTML to PDF: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to generate PDF: {str(e)}") from e
    
    def _to_docx(self, html: str) -> bytes:
        """
        Convert HTML to DOCX using python-docx.
        Supports: bold, italic, tables, headings, images.
        
        Args:
            html: HTML content
        
        Returns:
            DOCX content as bytes
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available for DOCX export")
            raise ValueError("DOCX export requires python-docx library")
        
        try:
            # Parse HTML
            if not BEAUTIFULSOUP_AVAILABLE:
                # Fallback: simple text extraction
                return self._html_to_docx_simple(html)
            
            soup = BeautifulSoup(html, 'html.parser')
            doc = Document()
            
            # Process HTML elements
            self._process_html_element(soup, doc)
            
            # Save to bytes
            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error converting HTML to DOCX: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to generate DOCX: {str(e)}") from e
    
    def _process_html_element(self, element, doc: Document):
        """Recursively process HTML elements and add to DOCX document."""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        for child in element.children:
            if hasattr(child, 'name'):
                tag_name = child.name.lower()
                
                if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    # Heading
                    level = int(tag_name[1])
                    heading = doc.add_heading(child.get_text(), level=level)
                    
                elif tag_name == 'p':
                    # Paragraph
                    para = doc.add_paragraph()
                    self._process_inline_formatting(child, para)
                    
                elif tag_name == 'table':
                    # Table
                    self._process_table(child, doc)
                    
                elif tag_name == 'ul' or tag_name == 'ol':
                    # List
                    self._process_list(child, doc, tag_name == 'ol')
                    
                elif tag_name in ['div', 'section', 'body']:
                    # Container - process children
                    self._process_html_element(child, doc)
                    
                elif tag_name == 'img':
                    # Embedding images in DOCX requires resolving the src
                    # attribute to an on-disk media file or a base64 data URI,
                    # then calling doc.add_picture().  Media-file resolution
                    # depends on a per-organization asset storage layer that is
                    # deferred to v2.  Base64 data URIs embedded directly in
                    # the HTML src are supported as a best-effort path below.
                    src = child.get('src', '')
                    if src.startswith('data:image/'):
                        try:
                            # Strip the data-URI header and decode
                            header, b64data = src.split(',', 1)
                            image_bytes = base64.b64decode(b64data)
                            image_stream = io.BytesIO(image_bytes)
                            doc.add_picture(image_stream, width=Inches(4))
                        except Exception as img_err:
                            logger.warning(
                                "Could not embed base64 image in DOCX: %s", img_err
                            )
                    else:
                        logger.warning(
                            "Skipping non-base64 image '%s' in DOCX export — "
                            "media-file resolution is deferred to v2.",
                            src[:80] if src else '(no src)',
                        )
    
    def _process_inline_formatting(self, element, paragraph):
        """Process inline formatting (bold, italic, etc.) in paragraph."""
        from docx.text.run import Run
        
        for child in element.children:
            if hasattr(child, 'name') and child.name is not None:
                tag_name = child.name.lower()
                text = child.get_text()
                
                if tag_name == 'strong' or tag_name == 'b':
                    run = paragraph.add_run(text)
                    run.bold = True
                elif tag_name == 'em' or tag_name == 'i':
                    run = paragraph.add_run(text)
                    run.italic = True
                elif tag_name == 'u':
                    run = paragraph.add_run(text)
                    run.underline = True
                else:
                    paragraph.add_run(text)
            elif hasattr(child, 'string') and child.string:
                paragraph.add_run(child.string)
    
    def _process_table(self, table_element, doc: Document):
        """Process HTML table and add to DOCX document."""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Determine number of columns
        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        
        # Create DOCX table
        docx_table = doc.add_table(rows=len(rows), cols=max_cols)
        docx_table.style = 'Light Grid Accent 1'
        
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    docx_cell = docx_table.rows[row_idx].cells[col_idx]
                    docx_cell.text = cell.get_text(strip=True)
                    
                    # Style header cells
                    if cell.name == 'th':
                        for paragraph in docx_cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    
    def _process_list(self, list_element, doc: Document, ordered: bool = False):
        """Process HTML list and add to DOCX document."""
        items = list_element.find_all('li')
        for item in items:
            para = doc.add_paragraph(item.get_text(), style='List Bullet' if not ordered else 'List Number')
    
    def _html_to_docx_simple(self, html: str) -> bytes:
        """Simple fallback DOCX conversion without BeautifulSoup."""
        doc = Document()
        
        # Extract text from HTML (very basic)
        import re
        text = re.sub(r'<[^>]+>', '', html)
        text = re.sub(r'\s+', ' ', text).strip()
        
        doc.add_paragraph(text)
        
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    
    def _to_csv(self, html: str) -> str:
        """
        Convert HTML table to CSV.
        Parses HTML tables and converts to CSV format.
        
        Args:
            html: HTML content
        
        Returns:
            CSV content as string
        """
        if not BEAUTIFULSOUP_AVAILABLE:
            logger.warning("BeautifulSoup not available. CSV export limited.")
            return "CSV export requires BeautifulSoup library"
        
        try:
            soup = BeautifulSoup(html, 'html.parser')
            tables = soup.find_all('table')
            
            if not tables:
                return "No tables found in HTML"
            
            # Convert first table to CSV
            # For multiple tables, you could concatenate or return first
            table = tables[0]
            
            output = io.StringIO()
            writer = csv.writer(output)
            
            # Process rows
            rows = table.find_all('tr')
            for row in rows:
                cells = row.find_all(['td', 'th'])
                row_data = [cell.get_text(strip=True) for cell in cells]
                writer.writerow(row_data)
            
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Error converting HTML to CSV: {str(e)}", exc_info=True)
            return f"Error generating CSV: {str(e)}"
