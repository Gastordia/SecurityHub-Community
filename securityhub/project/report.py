import io
import logging
import urllib
import os
import requests
from copy import deepcopy
import bleach
import pygal
from django.conf import settings
from django.db.models import Q
from django.http import HttpResponse
from rest_framework.response import Response
from django.template.loader import render_to_string
from django.template import TemplateDoesNotExist
from pygal.style import Style
# WeasyPrint import is lazy-loaded to avoid startup errors if GTK+ libraries are missing
from docxtpl import DocxTemplate,RichText
from docx import Document
from docx.table import Table
from utils.validators import is_safe_fetch_url
from utils.input_validation import sanitize_filename
from django.shortcuts import get_object_or_404
from datetime import datetime
from jinja2.sandbox import SandboxedEnvironment
import html
import traceback


from accounts.models import CustomUser
from configapi.bundled_docx_templates import get_bundled_docx_template
from .models import ProjectScope, Project, Vulnerability, VulnerableInstance
from configapi.models import ReportTemplate
from utils.doc_style import get_subdoc ,main_doc_style
from .services.project_contacts import get_project_manager_queryset
logger = logging.getLogger(__name__)

# Lazy import WeasyPrint to avoid startup errors if GTK+ libraries are missing
WEASYPRINT_AVAILABLE = False
WeasyHTML = None
try:
    from weasyprint import HTML as WeasyHTML
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError) as e:
    logger.warning(f"WeasyPrint not available: {e}. PDF report generation will be disabled. "
                   f"Install GTK+ runtime libraries to enable PDF reports. "
                   f"See: https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation")
    WEASYPRINT_AVAILABLE = False

# Global variable to store base_url
base_url = ""


def _plain_text(value):
    return html.unescape(bleach.clean(value or "", tags=[], strip=True)).strip()


def _company_name_for_project(project):
    organization = getattr(project, "organization", None)
    if organization and getattr(organization, "name", None):
        return organization.name
    return ""


def _replace_text_in_paragraph(paragraph, replacements):
    text = paragraph.text
    updated = text
    for old, new in replacements.items():
        updated = updated.replace(old, new)
    if updated == text:
        return
    if paragraph.runs:
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(updated)


def _replace_text_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_text_in_paragraph(paragraph, replacements)


def _set_cell_text(cell, value):
    text = value if isinstance(value, str) else str(value)
    if cell.paragraphs:
        cell.paragraphs[0].text = text
        for paragraph in cell.paragraphs[1:]:
            paragraph.text = ""
    else:
        cell.text = text


def _remove_table(table):
    tbl = table._element
    parent = tbl.getparent()
    if parent is not None:
        parent.remove(tbl)


def _table_from_element(doc, element):
    for table in doc.tables:
        if table._element is element:
            return table
    return Table(element, doc._body)


def _severity_to_fr(severity):
    severity_key = (severity or "").strip().lower()
    if severity_key == "critical":
        return "CRITIQUE"
    if severity_key == "high":
        return "MAJEUR"
    if severity_key == "medium":
        return "MODERE"
    if severity_key == "low":
        return "MINEUR"
    return "INFORMATION"


def _severity_bullet_label(severity):
    return f"● {_severity_to_fr(severity)}"


def _status_to_fr(status_value):
    mapping = {
        "Vulnerable": "En cours",
        "Confirm Fixed": "Corrige",
        "Accepted Risk": "Accepte",
    }
    return mapping.get(status_value, status_value or "")


def _overall_risk_label(vulnerabilities):
    ordered = ["Critical", "High", "Medium", "Low"]
    for severity in ordered:
        if any((v.vulnerabilityseverity or "").lower() == severity.lower() for v in vulnerabilities):
            return {
                "Critical": "Faible",
                "High": "Moyen",
                "Medium": "Moyen",
                "Low": "Satisfaisant",
            }[severity]
    return "Satisfaisant"


def _scope_summary(projectscope):
    lines = []
    for scope in projectscope:
        line = scope.scope or ""
        if scope.description:
            line = f"{line} - {scope.description}"
        lines.append(line.strip(" -"))
    return "\n".join(filter(None, lines))


def _instances_summary(vulnerability):
    instances = getattr(vulnerability, "instances_data", None)
    if not instances:
        return "Voir le perimetre du projet."
    lines = []
    for instance in instances:
        parts = [instance.get("URL", "").strip(), instance.get("Parameter", "").strip()]
        line = " | ".join([part for part in parts if part])
        if line:
            lines.append(line)
    return "\n".join(lines) or "Voir le perimetre du projet."


def _build_bundled_template_context(project, vulnerabilities, report_type, projectscope):
    return {
        "client_name": _company_name_for_project(project),
        "report_date": datetime.now().strftime("%d/%m/%Y"),
        "report_date_long": datetime.now().strftime("%B %d, %Y"),
        "scope_summary": _scope_summary(projectscope),
        "overall_risk": _overall_risk_label(vulnerabilities),
        "report_type": report_type,
    }


def _render_bundled_docx_template(template_info, project, vulnerabilities, report_type, projectscope):
    doc = Document(template_info["path"])
    context = _build_bundled_template_context(project, vulnerabilities, report_type, projectscope)

    replacements = {
        "[Nom du client]": context["client_name"] or "Client",
        "[JJ/MM/AAAA]": context["report_date"],
        "[Date]": context["report_date"],
        "[NIVEAU : Faible / Moyen / Satisfaisant]": context["overall_risk"],
        "Xx.xx.xx.xx   /   [plage IP / URLs testées]": context["scope_summary"] or "Perimetre a renseigner",
        "Xx.xx.xx.xx  /  [ plage réseau testée]": context["scope_summary"] or "Perimetre a renseigner",
        "[plage IP / URLs testées]": context["scope_summary"] or "Perimetre a renseigner",
        "[ plage réseau testée]": context["scope_summary"] or "Perimetre a renseigner",
    }

    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        _replace_text_in_table(table, replacements)

    if len(doc.tables) < 7:
        raise ValueError("Bundled DOCX template structure is invalid")

    summary_table = doc.tables[5]
    while len(summary_table.rows) > 1:
        summary_table._tbl.remove(summary_table.rows[-1]._tr)
    prefix = "E" if template_info["variant"] == "external" else "I"
    for index, vulnerability in enumerate(vulnerabilities, start=1):
        row = summary_table.add_row().cells
        row[0].text = f"{prefix}-{index:03d}"
        row[1].text = _severity_to_fr(vulnerability.vulnerabilityseverity)
        row[2].text = vulnerability.vulnerabilityname or "Vulnerabilite"
        row[3].text = _status_to_fr(vulnerability.status)

    detail_template = doc.tables[6]
    detail_tables = doc.tables[6:]

    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith(("E - ", "I-", "I- ", "E-")) and ("Missing Headers" in text or "[Titre" in text):
            p = paragraph._element
            parent = p.getparent()
            if parent is not None:
                parent.remove(p)

    anchor = None
    for index, vulnerability in enumerate(vulnerabilities, start=1):
        if index <= len(detail_tables):
            table = detail_tables[index - 1]
            anchor = table._element
        else:
            cloned_tbl = deepcopy(detail_template._element)
            insert_after = anchor if anchor is not None else detail_tables[-1]._element
            insert_after.addnext(cloned_tbl)
            table = _table_from_element(doc, cloned_tbl)
            anchor = cloned_tbl
        cells = table.rows[0].cells
        _set_cell_text(cells[0], f"{prefix}-{index:03d}")
        _set_cell_text(cells[1], vulnerability.vulnerabilityname or "Vulnerabilite")
        _set_cell_text(cells[2], _severity_bullet_label(vulnerability.vulnerabilityseverity))

        impacted = _instances_summary(vulnerability)
        _set_cell_text(table.rows[1].cells[1], impacted)
        _set_cell_text(table.rows[1].cells[2], impacted)

        description_parts = [_plain_text(vulnerability.vulnerabilitydescription)]
        poc_text = _plain_text(vulnerability.POC)
        if poc_text:
            description_parts.append(f"Preuve: {poc_text}")
        description = "\n\n".join([part for part in description_parts if part]) or "Voir details dans SecurityHub."
        _set_cell_text(table.rows[2].cells[1], description)
        _set_cell_text(table.rows[2].cells[2], description)

        recommendation_parts = [_plain_text(vulnerability.vulnerabilitysolution)]
        references_text = _plain_text(vulnerability.vulnerabilityreferlnk)
        if references_text:
            recommendation_parts.append(f"References: {references_text}")
        recommendation = "\n\n".join([part for part in recommendation_parts if part]) or "Aucune recommandation renseignee."
        _set_cell_text(table.rows[3].cells[1], recommendation)
        _set_cell_text(table.rows[3].cells[2], recommendation)

    for table in detail_tables[len(vulnerabilities):]:
        _remove_table(table)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    response = HttpResponse(
        output.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    safe_name = sanitize_filename(project.name or 'report')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}_vulnerability_report.docx"'
    return response


def CheckReport(Report_format,Report_type,pk,url,standard,request,access_token,is_staff,template_id=None,user=None):
    """
    Generate report with optional template support.

    Args:
        Report_format: Format of report (pdf, docx, excel)
        Report_type: Type of report (Audit, Re-Audit)
        pk: Project ID
        url: Base URL
        standard: Report standard
        request: Django request object
        access_token: Access token for images
        is_staff: Whether user is staff
        template_id: Optional template ID from database
        user: User object for tracking template usage
    """
    global base_url
    base_url = url if url is not None else ''

    if Report_format == "pdf":
        return GetHTML(
            Report_type, pk, standard, request, is_staff, access_token=access_token, template_id=template_id, user=user
        )
    if Report_format == "docx":
        return generate_vulnerability_document(
            pk, Report_type, standard, access_token=access_token, template_id=template_id, user=user
        )
    return Response(
        {"Status": "Failed", "Message": f"Unsupported report format: {Report_format!r}"},
        status=400,
    )



def generate_vulnerability_document(pk,Report_type,standard,access_token=None,template_id=None,user=None):
    try:
        project_id = pk
        project = get_object_or_404(Project, id=project_id)
        owners = project.owner.all()
        vuln = Vulnerability.objects.filter(project=project).order_by('-cvssscore')
        totalvulnerability = vuln.filter(project=project).count()
        totalretests_queryset = []

        projectscope = ProjectScope.objects.filter(project=project)
        bundled_template = get_bundled_docx_template(template_id)
        if bundled_template:
            project_description = _plain_text(project.description)
            project_exception = _plain_text(project.projectexception)
            for vulnerability in vuln:
                vulnerability.instances_data = [
                    {
                        'URL': instance.URL,
                        'Parameter': instance.Parameter if instance.Parameter is not None else '',
                        'Status': instance.status
                    }
                    for instance in VulnerableInstance.objects.filter(vulnerabilityid=vulnerability, project=project)
                    if instance.URL
                ]
            return _render_bundled_docx_template(
                bundled_template,
                project,
                list(vuln),
                Report_type,
                list(projectscope),
            )
        
        # Use database template if provided, otherwise use default
        if template_id:
            try:
                from configapi.services.template_service import TemplateService
                
                template_service = TemplateService()
                # For DOCX, we still need to use the file-based template system
                # but we can track usage through the service
                # Note: Full DOCX template support would require storing actual .docx files
                template_path = os.path.join(settings.BASE_DIR, 'templates', 'report.docx')
                
                # Track template usage
                try:
                    db_template = ReportTemplate.objects.get(
                        id=template_id,
                        format__in=['docx', 'word'],
                        is_active=True
                    )
                    template_service._track_usage(
                        template=db_template,
                        format='docx',
                        user=user,
                        project=project
                    )
                except ReportTemplate.DoesNotExist:
                    logger.warning(f"Template {template_id} not found or not suitable for DOCX, using default")
            except Exception as e:
                logger.warning(f"Error tracking template usage for DOCX: {str(e)}")
                template_path = os.path.join(settings.BASE_DIR, 'templates', 'report.docx')
        else:
            template_path = os.path.join(settings.BASE_DIR, 'templates', 'report.docx')
        
        doc = DocxTemplate(template_path)
        projectmanagers = get_project_manager_queryset(project)
        customeruser = CustomUser.objects.filter(is_active=True)
        mycomany = None
        project_description = get_subdoc(doc, project.description, access_token, base_url)
        project_exception = get_subdoc(doc, project.projectexception, access_token, base_url)

        for vulnerability in vuln:
            # Convert CKEditor fields from HTML to DOCX format
            vulnerability.vulnerabilitydescription = get_subdoc(doc, vulnerability.vulnerabilitydescription, access_token, base_url)
            vulnerability.POC = get_subdoc(doc, vulnerability.POC, access_token, base_url)
            vulnerability.vulnerabilitysolution = get_subdoc(doc, vulnerability.vulnerabilitysolution, access_token, base_url)
            vulnerability.vulnerabilityreferlnk = get_subdoc(doc, vulnerability.vulnerabilityreferlnk, access_token, base_url)

            vulnerability.instances_data = [
                {
                    'URL': instance.URL,
                    'Parameter': instance.Parameter if instance.Parameter is not None else '',  # Set empty string for None
                    'Status': instance.status
                }
                for instance in VulnerableInstance.objects.filter(vulnerabilityid=vulnerability, project=project)
                if instance.URL  # Exclude instances with empty URL
            ]

        totalretest = []
        currentdate=datetime.now()
        context = {'project': project, 'vulnerabilities': vuln,'Report_type':Report_type,'mycomany':mycomany,'projectmanagers':projectmanagers,'customeruser':customeruser,'owners': owners,
                'project_exception':project_exception,'project_description':project_description,'currentdate':currentdate,
                'standard':standard,'totalvulnerability':totalvulnerability,'totalretest':totalretest,'projectscope':projectscope,
                'page_break': RichText('\f'),'new_line': RichText('\n')
                }
        jinja_env = SandboxedEnvironment(autoescape=True)
        jinja_env.trim_blocks = True
        jinja_env.lstrip_blocks = True
        doc.render(context,jinja_env)
        doc = main_doc_style(doc)
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        safe_name = sanitize_filename(project.name or 'report')
        response['Content-Disposition'] = f'attachment; filename="{safe_name}_vulnerability_report.docx"'
        doc.save(response)
        return response

    except Exception as e:
        logger.error(f"An error occurred: {e}")
        logger.error("Traceback: " + traceback.format_exc())
        return Response({"Status": "Failed", "Message": "Something Went Wrong"})


def GetHTML(Report_type,pk,standard,request,is_staff,access_token=None,template_id=None,user=None):
    """
    Generate PDF report from HTML template.
    
    Args:
        Report_type: Type of report (Audit, Re-Audit)
        pk: Project ID
        standard: Report standard
        request: Django request object
        is_staff: Whether user is staff
        template_id: Optional template ID from database
        user: User object for tracking template usage
    """
    try:
        if standard is None:
            standard = []

        project = Project.objects.get(pk=pk)        ## Get All Projects Vulnerability filter higher to lower CVSS Score
        vuln = Vulnerability.objects.filter(project=project).order_by('-cvssscore')
        instances = VulnerableInstance.objects.filter(project=project)
        projectmanagers = get_project_manager_queryset(project)
        customeruser = CustomUser.objects.filter(is_active=True)


        ciritcal =  vuln.filter(project=project,vulnerabilityseverity='Critical',status='Vulnerable').count()
        high =  vuln.filter(project=project,vulnerabilityseverity='High',status='Vulnerable').count()
        medium =  vuln.filter(project=project,vulnerabilityseverity='Medium',status='Vulnerable').count()
        low =  vuln.filter(project=project,vulnerabilityseverity='Low',status='Vulnerable').count()
        info = vuln.filter((Q(status='Vulnerable')) & (Q(vulnerabilityseverity='Informational') | Q(vulnerabilityseverity='None'))).count()

        custom_style = Style(
            colors=("#FF491C", "#F66E09", "#FBBC02", "#20B803", "#3399FF"),
            background='transparent',
            plot_background='transparent',
            legend_font_size=0,
            legend_box_size=0,
            value_font_size=40
            )
        pie_chart = pygal.Pie(style=custom_style)
        pie_chart.legend_box_size = 0

        pie_chart.add('Critical', ciritcal)
        pie_chart.add('High', high)
        pie_chart.add('Medium', medium)
        pie_chart.add('Low', low)
        pie_chart.add('Informational', info)

        pie_rendered = pie_chart.render(is_unicode=True)
        if pie_rendered is None:
            pie_rendered = ''
        elif isinstance(pie_rendered, bytes):
            pie_rendered = pie_rendered.decode('utf-8', errors='replace')

        ### Get Total Vulnerability Count
        totalvulnerability = vuln.filter(project=project).count()
        mycomany = None


        ### Get All Scope from the project
        projectscope = ProjectScope.objects.filter(project=project)

        totalretest = []
        data = {
            'projectscope': projectscope,
            'totalvulnerability': totalvulnerability,
            'standard': standard,
            'Report_type': Report_type,
            'mycomany': mycomany,
            'totalretest': totalretest,
            'vuln': vuln,
            'project': project,
            'ciritcal': ciritcal,
            'high': high,
            'medium': medium,
            'low': low,
            'info': info,
            'instances': instances,
            'projectmanagers': projectmanagers,
            'customeruser': customeruser,
            'internalusers': [],
            'pie_chart': pie_rendered,
        }
        try:
            # Use database template if provided, otherwise use default
            if template_id:
                try:
                    from configapi.services.template_engine import (
                        TemplateEngine, TemplateNotFoundError, TemplateEngineError
                    )
                    
                    engine = TemplateEngine()
                    rendered_content = engine.render(
                        template_id=template_id,
                        report_data=data,
                        format='html',  # PDF is generated from HTML
                        validate=True,
                        track_usage=True,
                        user=user,
                        project=project,
                        base_url=base_url
                    )
                    # Ensure string output for HTML
                    if isinstance(rendered_content, bytes):
                        rendered_content = rendered_content.decode('utf-8')
                except TemplateNotFoundError:
                    logger.warning(f"Template {template_id} not found or not suitable for PDF, using default")
                    rendered_content = render_to_string('report.html', data, request=request)
                except TemplateEngineError as e:
                    logger.error(f"Error rendering template {template_id}: {str(e)}")
                    # Fallback to default template on render error
                    rendered_content = render_to_string('report.html', data, request=request)
                except Exception as e:
                    logger.error(f"Unexpected error with template {template_id}: {str(e)}", exc_info=True)
                    # Fallback to default template on unexpected error
                    rendered_content = render_to_string('report.html', data, request=request)
            else:
                # Use default template
                rendered_content = render_to_string('report.html', data, request=request)

            if rendered_content is None:
                logger.error("Report HTML rendering produced None; using empty document shell")
                rendered_content = '<html><body><p>Report content could not be rendered.</p></body></html>'
            elif isinstance(rendered_content, bytes):
                rendered_content = rendered_content.decode('utf-8', errors='replace')
            
            response = generate_pdf_report(rendered_content, base_url, access_token=access_token)
            safe_name = sanitize_filename(project.name or 'report')
            response['Content-Disposition'] = f'attachment; filename="{safe_name}_vulnerability_report.pdf"'
            return response
        except (TemplateDoesNotExist, IOError) as e:
            # Handle template not found error
            logger.error("Template or file error: %s", e, exc_info=True)
            return Response({"Status": "Failed", "Message": f"Template or file error: {str(e)}"})
        except Exception as e:
            # Handle any other errors
            logger.error(f"Unexpected error in GetHTML: {str(e)}", exc_info=True)
            return Response({"Status": "Failed", "Message": f"Error generating report: {str(e)}"})

    except Exception as e:
        logger.error(f"An error occurred in GetHTML: {str(e)}", exc_info=True)
        return Response({"Status": "Failed", "Message": f"Error generating report: {str(e)}"})


def generate_pdf_report(rendered_content, base_url, access_token=None):
    """
    Generate PDF report from rendered HTML content.

    Args:
        rendered_content: Rendered HTML content
        base_url: Base URL for resolving relative paths
        access_token: JWT access token for authenticated image fetching

    Returns:
        HttpResponse with PDF content or error response
    """
    if not WEASYPRINT_AVAILABLE:
        error_msg = (
            "PDF report generation is not available. WeasyPrint requires GTK+ runtime libraries. "
            "Please install GTK+ runtime for Windows from: "
            "https://github.com/tschoonj/GTK-for-Windows-Runtime-Environment-Installer/releases"
        )
        logger.error(error_msg)
        return Response({"Status": "Failed", "Message": error_msg})

    if rendered_content is None:
        logger.error("generate_pdf_report called with rendered_content=None")
        return Response(
            {
                "Status": "Failed",
                "Message": "Report HTML was empty; cannot build PDF. Check project data and templates.",
            }
        )
    if isinstance(rendered_content, bytes):
        rendered_content = rendered_content.decode('utf-8', errors='replace')
    if not isinstance(rendered_content, str):
        rendered_content = str(rendered_content)

    fetcher = make_fetcher(access_token)
    safe_base = base_url if isinstance(base_url, str) and base_url.strip() else "file:///"
    try:
        pdf = WeasyHTML(
            string=rendered_content,
            url_fetcher=fetcher,
            base_url=safe_base,
        ).write_pdf()
        # Return the PDF response
        response = HttpResponse(pdf, content_type='application/pdf')
        return response
    except Exception as e:
        # Return a server error response if there's an issue
        logger.error("Error generating PDF report: %s", e, exc_info=True)
        return Response({"Status": "Failed", "Message": f"Failed to generate PDF: {str(e)}"})


def is_whitelisted(url):
    """Checks if the given URL is whitelisted to protect against SSRF to access internal or external network."""
    if is_safe_fetch_url(url):
        return True
    logger.error("URL is not Whitelisted Check the %s", url)
    return False



def make_fetcher(access_token):
    """Return a WeasyPrint url_fetcher closure that carries the JWT access token."""
    def my_fetcher(url):
        # WeasyPrint may call the fetcher with None; delegate to default handling
        if url is None:
            return None
        if not isinstance(url, str):
            url = str(url)
        url = url.strip()
        if not url:
            return None

        # Check if the URL is whitelisted
        if is_whitelisted(url):
            # allow_redirects=False prevents redirect-based SSRF to internal resources
            if "/api/project/getimage/" in url:
                headers = {"Authorization": f"Bearer {access_token}"}
                response = requests.get(url, headers=headers, verify=True,
                                        timeout=10, allow_redirects=False)
                response.raise_for_status()
                enc = response.encoding or "utf-8"
                return {
                    "string": response.content,
                    "mime_type": response.headers.get("Content-Type", "image/jpeg") or "image/jpeg",
                    "encoding": enc,
                    "redirected_url": url,
                }
            else:
                response = requests.get(url, verify=True, timeout=10, allow_redirects=False)
                response.raise_for_status()
                mime_type = response.headers.get("Content-Type", "application/octet-stream") or "application/octet-stream"
                enc = response.encoding or "utf-8"
                return {
                    "string": response.content,
                    "mime_type": mime_type,
                    "encoding": enc,
                    "redirected_url": url,
                }
        else:
            raise ValueError(f'URL is Not WhiteListed for: {url!r}')
    return my_fetcher
